import streamlit as st
import google.generativeai as genai
from PIL import Image
import docx
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import PyPDF2
import io
import os
import time
from datetime import datetime

# ==========================================
# 0. 自动版本号生成逻辑
# ==========================================
def get_app_version():
    try:
        timestamp = os.path.getmtime(__file__)
        dt = datetime.fromtimestamp(timestamp)
        build_ver = dt.strftime('%m%d.%H%M')
        return f"v13.34.{build_ver}", dt.strftime('%Y-%m-%d %H:%M:%S')
    except Exception:
        return "v13.34.Dev", "Unknown"

current_version, last_updated_time = get_app_version()

# ==========================================
# 1. 页面基础配置
# ==========================================
st.set_page_config(page_title="个人陈述写作", layout="wide")

# ==========================================
# UI 样式注入
# ==========================================
def apply_custom_css():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    
    :root {
        --primary-color: #3666FA;
        --bg-color: #FBF7EC;
        --text-color: #3666FA;
        --button-text: #FBF7EC;
    }

    html, body, [class*="css"] {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
        color: var(--text-color);
        background-color: var(--bg-color);
    }
    
    header {visibility: hidden;}
    footer {visibility: hidden;}

    .stApp {
        background-color: var(--bg-color);
    }

    [data-testid="stSidebar"] {
        background-color: #0f172a;
        border-right: 1px solid #1e293b;
    }
    
    [data-testid="stSidebar"] h1,
    [data-testid="stSidebar"] h2, 
    [data-testid="stSidebar"] h3, 
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] span,
    [data-testid="stSidebar"] .stMarkdown,
    [data-testid="stSidebar"] div {
        color: #ffffff !important;
    }
    
    [data-testid="stSidebar"] hr {
        border-color: #334155 !important;
    }

    [data-testid="stSidebar"] .stTextInput input {
        background-color: #1e293b !important; 
        color: #ffffff !important;
        border: 1px solid #334155 !important;
    }

    [data-testid="stSidebar"] .stSelectbox div[data-baseweb="select"] {
        background-color: #1e293b !important; 
        border: 1px solid #334155 !important; 
    }
    
    [data-testid="stSidebar"] .stSelectbox div[data-baseweb="select"] * {
        color: #1e293b !important;                    
        font-family: 'Inter', sans-serif !important;  
    }
    
    [data-testid="stSidebar"] .stSelectbox svg {
        fill: #ffffff !important;
    }

    h1 {
        color: var(--text-color) !important;
        font-weight: 800 !important;
        font-size: 2.5rem !important;
        letter-spacing: -0.02em;
        margin-bottom: 2rem !important;
        text-align: left !important;
    }
    
    h2, h3 {
        color: var(--text-color) !important;
        font-weight: 600 !important;
        margin-top: 1rem !important;
        margin-bottom: 1rem !important;
    }
    
    .main p, .main label, .main .stMarkdown, .main .stText, .main .stCaption {
        color: var(--text-color) !important;
    }

    .main .stTextInput input, .main .stTextArea textarea, .main .stSelectbox div[data-baseweb="select"] {
        border: 1px solid rgba(54, 102, 250, 0.3) !important;
        border-radius: 8px !important;
        padding: 0.6rem 0.8rem !important;
        background-color: #ffffff !important;
        font-size: 13px !important;
        color: #1e293b !important;
        transition: all 0.2s ease;
    }

    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: var(--primary-color) !important;
        box-shadow: 0 0 0 2px rgba(54, 102, 250, 0.1) !important;
    }

    div.stButton > button[kind="primary"] {
        background-color: var(--primary-color) !important;
        color: #ffffff !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 0.6rem 1.5rem !important;
        font-weight: 500 !important;
        box-shadow: 0 1px 2px rgba(54, 102, 250, 0.2) !important;
        transition: all 0.2s ease !important;
        width: 100%; 
    }

    div.stButton > button[kind="secondary"] {
        background-color: #E2E8F0 !important;
        color: #64748B !important;            
        border: none !important;
        border-radius: 8px !important;
        padding: 0.6rem 1.5rem !important;
        font-weight: 500 !important;
        width: 100%; 
    }

    div.stButton > button:hover {
        opacity: 0.9;
        transform: translateY(-1px);
    }
    
    .stDownloadButton button {
        background-color: var(--primary-color) !important;
        color: var(--button-text) !important;
        border: none !important;
    }

    .streamlit-expanderHeader {
        background-color: #ffffff !important;
        border: 1px solid rgba(54, 102, 250, 0.2) !important;
        border-radius: 8px !important;
        color: var(--text-color) !important;
        font-weight: 600 !important;
    }
    
    [data-testid="stFileUploader"] {
        border: 1px dashed rgba(54, 102, 250, 0.4);
        background-color: #ffffff;
        border-radius: 8px;
        padding: 1rem;
        padding-bottom: 20px;
    }
    [data-testid="stFileUploader"]:hover {
        border-color: var(--primary-color);
        background-color: rgba(54, 102, 250, 0.05);
    }

    .block-container {
        padding-top: 3rem !important;
        padding-bottom: 3rem !important;
        max-width: 1200px !important;
    }
    
    hr {
        border-color: rgba(54, 102, 250, 0.2) !important;
    }
    
    .stAlert {
        background-color: #1e293b !important;
        border: none !important;
        color: #ffffff !important;
    }

    .stRadio p {
        font-size: 13px !important;
    }

    div[data-testid="stHorizontalBlock"] {
        align-items: stretch;
        height: auto;
    }
    div[data-testid="column"] {
        display: flex;
        flex-direction: column;
        height: 100%;
    }
    div[data-testid="stVerticalBlockBorderWrapper"] {
        flex: 1 1 auto;
        height: 100%;
        display: flex;
        flex-direction: column;
        min-height: 450px;
        border-color: rgba(54, 102, 250, 0.2) !important;
        background-color: #ffffff !important;
    }
    div[data-testid="stVerticalBlockBorderWrapper"] > div {
        flex-grow: 1;
        display: flex;
        flex-direction: column;
    }
    .stMarkdown p {
        margin-bottom: 0px;
    }
    </style>
    """, unsafe_allow_html=True)

apply_custom_css()

# 初始化 Session State
if 'generated_sections' not in st.session_state:
    st.session_state['generated_sections'] = {}
if 'motivation_trends' not in st.session_state:
    st.session_state['motivation_trends'] = ""
if 'full_chinese_draft' not in st.session_state:
    st.session_state['full_chinese_draft'] = ""
if 'full_translated_text' not in st.session_state:
    st.session_state['full_translated_text'] = ""
if 'module_states' not in st.session_state:
    display_order = ["Motivation", "Academic", "Internship", "Why_School", "Career_Goal"]
    st.session_state['module_states'] = {key: True for key in display_order}

# 标题
st.title("个人陈述写作")
st.markdown("---")

# ==========================================
# 2. 核心辅助函数
# ==========================================

def render_blue_box(text):
    if "</div>" in text or "</ul>" in text:
        html_text = text
    else:
        html_text = text.replace('\n', '<br>')
        
    st.markdown(f"""
    <div style="
        background-color: #3666FA; 
        color: #ffffff; 
        padding: 15px 20px; 
        border-radius: 12px; 
        margin-bottom: 20px; 
        font-size: 13px; 
        line-height: 1.6;
        box-shadow: 0 2px 5px rgba(54, 102, 250, 0.2);
    ">
        {html_text}
    </div>
    """, unsafe_allow_html=True)

def set_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_word_docx(content, header_text, font_name, is_chinese=False):
    doc = docx.Document()
    
    section = doc.sections[0]
    header = section.header
    
    header_para = header.paragraphs[0]
    header_para.text = header_text
    header_para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT 
    
    set_bottom_border(header_para)
    
    for run in header_para.runs:
        run.font.name = font_name
        run.font.size = Pt(12)
        run.font.italic = True
        if is_chinese:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            
    content = content.replace("**", "")
    content = content.replace("*", "")
    
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith("---") and line.endswith("---"):
            continue
            
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = font_name
            run.font.size = Pt(11)
            if is_chinese:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# 3. 系统设置
# ==========================================
with st.sidebar:
    st.header("系统设置")
    
    api_key = st.text_input("请输入 Google API Key", type="password", help="请在 Google AI Studio 申请 Key")
    
    if not api_key:
        st.warning("请输入 Key")
    else:
        st.success("Key 已就绪")
    
    model_name = st.selectbox("选择模型", ["gemini-3-pro-preview", "gemini-2.5-pro"], index=0)

# ==========================================
# 4. 核心函数 (优化效率与稳健性)
# ==========================================

# 效率优化：缓存模型加载，避免重复初始化
@st.cache_resource
def load_model(api_key, model_name):
    genai.configure(api_key=api_key)
    return genai.GenerativeModel(model_name)

# 效率优化：压缩图片，防止上传过大图片导致 Token 消耗过大
def process_image(img_file):
    image = Image.open(img_file)
    max_size = 1024
    if max(image.size) > max_size:
        image.thumbnail((max_size, max_size))
    return image

def read_word_file(file):
    try:
        doc = docx.Document(file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        return f"[读取 Word 失败]: {e}"

def read_pdf_text(file):
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text if text else "[PDF 内容为空，请检查文件]"
    except Exception as e:
        return f"[读取 PDF 失败]: {e}"

def get_gemini_response(prompt, media_content=None, text_context=None):
    if not api_key:
        return "错误: 请先在左侧侧边栏输入 API Key"
        
    try:
        # 使用缓存的模型加载函数
        model = load_model(api_key, model_name)
        
        content = []
        content.append(prompt)
        
        if text_context:
            content.append(f"\n【参考文档/背景信息】:\n{text_context}")
        
        if media_content:
            if isinstance(media_content, list):
                content.extend(media_content)
            else:
                content.append(media_content)
        
        response = model.generate_content(content)
        return response.text
    except Exception as e:
        return f"API 调用错误: {str(e)}"

# ==========================================
# 5. 界面：信息采集
# ==========================================
st.header("信息采集与素材上传")

col_student, col_counselor, col_target = st.columns(3)

with col_student:
    with st.container(border=True):
        st.markdown("### 学生提供信息")
        st.caption("上传简历、素材表与成绩单")
        
        uploaded_material = st.file_uploader("文书素材/简历 (Word/PDF)", type=['docx', 'pdf'])
        uploaded_transcript = st.file_uploader("成绩单 (截图/PDF)", type=['png', 'jpg', 'jpeg', 'pdf'])

with col_counselor:
    with st.container(border=True):
        st.markdown("### 顾问指导意见")
        st.caption("设定文书的整体策略与调性")
        
        counselor_strategy = st.text_area(
            "写作策略/人设强调", 
            height=280, 
            placeholder="例如：\n1. 强调量化背景\n2. 解释GPA劣势\n3. 突出某段实习的领导力..."
        )

with col_target:
    with st.container(border=True):
        st.markdown("### 目标专业信息")
        st.caption("输入目标学校与课程设置")
        
        target_school_name = st.text_input("目标学校 & 专业", placeholder="例如：UCL - MSc Business Analytics")
        
        st.markdown("**课程设置**") 
        
        tab_text, tab_img = st.tabs(["文本粘贴", "图片上传"])
        
        with tab_text:
            target_curriculum_text = st.text_area("粘贴课程列表", height=140, placeholder="Core Modules: ...", label_visibility="collapsed")
        
        with tab_img:
            uploaded_curriculum_images = st.file_uploader("上传课程截图", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True, label_visibility="collapsed")

student_background_text = ""
if uploaded_material:
    if uploaded_material.name.endswith('.docx'):
        student_background_text = read_word_file(uploaded_material)
    elif uploaded_material.name.endswith('.pdf'):
        student_background_text = read_pdf_text(uploaded_material)

# ==========================================
# 6. 界面：写作设定
# ==========================================
st.markdown("---")
st.header("写作设定")

display_order = ["Motivation", "Academic", "Internship", "Why_School", "Career_Goal"]
modules = {
    "Motivation": "申请动机",
    "Academic": "本科学习",
    "Internship": "实习/工作",
    "Why_School": "选校理由",
    "Career_Goal": "职业规划"
}

english_modules = {
    "Motivation": "Motivation",
    "Academic": "Academic Background",
    "Internship": "Professional Experience",
    "Why_School": "Why School",
    "Career_Goal": "Career Goal"
}

col_modules, col_style = st.columns([3, 1])

with col_modules:
    st.markdown("**选择模块 (点击切换选中状态):**")
    mod_cols = st.columns(len(display_order))
    
    for idx, key in enumerate(display_order):
        is_selected = st.session_state['module_states'][key]
        label = modules[key]
        btn_type = "primary" if is_selected else "secondary"
        
        if mod_cols[idx].button(label, key=f"btn_mod_{key}", type=btn_type, use_container_width=True):
            st.session_state['module_states'][key] = not st.session_state['module_states'][key]
            st.rerun()

    selected_modules = [key for key in display_order if st.session_state['module_states'][key]]

with col_style:
    spelling_preference = st.radio(
        "拼写偏好",
        ["英式 (British)", "美式 (American)"],
        help="翻译时将严格遵循所选的拼写习惯"
    )

# ==========================================
# 7. 核心逻辑：生成 Prompt
# ==========================================
st.markdown("---")
st.header("一键点击创作")

CLEAN_OUTPUT_RULES = """
【绝对输出规则】
1. 仅输出纯文本正文。
2. 禁止包含任何开场白、结束语或元说明。
3. 禁止使用任何Markdown格式（如标题、加粗、列表）。
4. 将所有内容合并成一个单一、连贯的中文自然段。
"""

TRANSLATION_RULES_BASE = """
将提供的中文文本翻译成一段专业的、语气自然的英文个人陈述。
一. 输出格式
整体结构：输出为单一的、完整的纯文本段落。
标点规则：所有逗号和句号必须置于闭合引号之外。
禁止格式：禁止使用任何Markdown。
二. 核心写作风格与规则
【句子结构与叙事流】
主语选择：优先使用人称主语。
“, verb-ing”结构：谨慎使用“, verb-ing”的句式。
分号使用：适度使用分号连接紧密相关的观点。
过渡语使用：尽量添加多样的过渡短语。
【词汇与语气】
简洁词汇：使用精准、简单的词语。
三. 绝对禁止项
禁止句式：严禁使用“I did X, thereby/thus/enabling me to do Y”结构。
禁止副词组合：完全禁止“副词+动词”和“副词+形容词”用法。
禁止词汇清单：严格禁止使用以下词汇: delve into, uncover, master, cultivate, address, bridge, spearhead, pioneer, align with, stems from, underscore, highlight, pivotal, seamless, systematically, rigorously, profoundly, deeply, acutely, keenly, comprehensively, perfectly, meticulously, proficiency, Additionally, paradigm, trajectory, aspirations, vision, landscape, tapestry, realm, foundation, tenure, testament, commitment, thereby, thus, in turn, "not only... but also", "Building on this", "rich tapestry", "testament to", "a wide array of", "my goal is to", "focus will be"
"""

if st.button("开始生成初稿", type="primary"):
    if not api_key:
        st.error("请先在左侧侧边栏输入有效的 Google API Key")
        st.stop()

    has_curriculum = target_curriculum_text or uploaded_curriculum_images
    
    if not uploaded_material or not uploaded_transcript or not has_curriculum:
        st.error("请确保：文书素材/简历、成绩单、目标课程信息 均已提供。")
        st.stop()
    
    if not selected_modules:
        st.warning("请至少选择一个写作模块。")
        st.stop()
    
    # 稳健性优化：处理图片时增加异常捕获和压缩
    transcript_content = []
    if uploaded_transcript.type == "application/pdf":
        transcript_content.append({
            "mime_type": "application/pdf",
            "data": uploaded_transcript.getvalue()
        })
    else:
        try:
            transcript_content.append(process_image(uploaded_transcript))
        except Exception as e:
            st.error(f"处理成绩单图片时出错: {e}")
            st.stop()

    curriculum_imgs = []
    if uploaded_curriculum_images:
        for img_file in uploaded_curriculum_images:
            try:
                curriculum_imgs.append(process_image(img_file))
            except Exception as e:
                st.warning(f"跳过一张损坏的课程图片: {e}")
    
    progress_bar = st.progress(0)
    status_container = st.empty()
    detail_container = st.empty()
    
    total_steps = len(selected_modules)
    current_step = 0
    
    st.session_state['generated_sections'] = {} 

    # Prompt 定义 (Update: Motivation Logic)
    prompt_motivation = f"""
    【任务】撰写 Personal Statement 的 "申请动机" (Motivation) 部分。
    
    【步骤 1：深度调研 (Research)】
    请先分析 {target_school_name} 所在领域的最新行业热点或学术趋势。
    **请严格列出 3 个关键趋势 (Options)**，并严格按照以下 **HTML 格式** 输出（除文献/报告标题保留原文外，其余分析内容请使用**中文**）：

    <div style="margin-bottom: 18px;">
        <div style="font-weight: bold; font-size: 14px; margin-bottom: 6px;">Option [X]: [Trend Title]</div>
        <ul style="margin: 0; padding-left: 18px; list-style-position: outside;">
            <li style="margin-bottom: 4px; line-height: 1.4;"><b>Source</b>: [Specific Paper Title/Report Name/News Source]</li>
            <li style="line-height: 1.4;"><b>Relevance</b>: [深度分析趋势与学生背景/项目的关联。解释为什么这个趋势对该学生重要，以及他们之前的经历（如具体项目、技能）如何与此契合。此部分必须详细展开。]</li>
        </ul>
    </div>

    【步骤 2：撰写正文 (Drafting)】
    **关键指令**：请从上述 3 个趋势中，**自主选择** 一个最能体现该专业价值的趋势作为核心切入点。
    
    **写作结构要求 (Strict Structure)**：
    1.  **开门见山 (The Hook)**：写 1-2 句精准、犀利的话。直接概括“我希望通过硕士学习继续探索 [具体细分领域]”或“获得 [具体高阶知识]”。
    2.  **行业洞察 (Industry Reflection)**：紧接着阐述你对该领域的深度思考。此处**必须**结合你选择的那个“调研趋势”进行议论，展示你对行业痛点或机遇的理解。
    3.  **锁定目标 (The Match)**：最后总结，明确表达“因此，我希望通过 {target_school_name} 的 [项目名称] 来实现这一目的，也即我的深造需求”。
    
    **负面清单 (Constraints)**：
    * **严禁**在这一段提及“我过去做了什么”、“我的本科经历”或“我的实习”。这些内容请留到后续段落，这里只谈未来目标和行业思考。
    * 不用出现具体的文献名称或来源引用，将趋势内化为自己的观点。
    
    【严格输出格式】
    请严格按照下方分隔符输出，不要包含其他内容：
    [TRENDS_START]
    (在此处列出 3 个调研趋势和来源，使用上述 HTML 格式)
    [TRENDS_END]
    [DRAFT_START]
    (在此处撰写正文段落，纯文本，无Markdown)
    [DRAFT_END]
    """

    prompt_career = f"""
    【任务】撰写 "职业规划" (Career Goals) 部分。
    【输入背景】
    - 目标专业: {target_school_name}
    - 顾问思路: {counselor_strategy}
    【内容要求】
    1. 核心视角：必须从应届毕业生的视角出发。
    2. 文本中必须明确提及至少一个具体的公司名称（例如：Google, 腾讯, etc.）。
    3. 文本中必须明确提及一个与该公司相匹配的具体的职位名称（例如：Data Scientist, Product Manager, etc.）。
    4. 必须将该职位的预期工作内容与个人未来的学习深造方向紧密结合。
    {CLEAN_OUTPUT_RULES}
    """

    prompt_academic = f"""
    【任务】撰写 "本科学习经历" (Academic Background) 部分。
    【输入背景】
    - 目标专业: {target_school_name}
    - 核心依据 (成绩单): 见附带文件 (PDF或图片)
    - 辅助参考 (学生素材/简历): 见附带文本
    【Core Principle：深度优先】
    仅精选2-3门与申请专业最相关的核心课程。
    可以提及课程名称，但严禁仅停留在名称层面。课程名必须作为学术反思的载体自然融入文段
    【内容要求 - 必须包含细节】
    1. 专业锚点：准确引出该课程涉及的核心模型、算法、理论或关键术语。
    2. 认知还原：结合具体素材，描述你对上述概念的独到理解、推导过程或实践应用。
    3. 价值闭环：论证该课程沉淀的学术能力如何支撑你在 {target_school_name} 的后续学习计划。
    4. 严禁使用 Bullet Points，必须是一段逻辑严密、首尾贯通的学术叙事（Academic Narrative）。
    {CLEAN_OUTPUT_RULES}
    """

    prompt_whyschool = f"""
    【任务】撰写 "Why School" 部分。
    【输入背景】
    - 目标学校: {target_school_name}
    - 顾问思路: {counselor_strategy}
    {f'【目标课程文本列表】:{target_curriculum_text}' if target_curriculum_text else ''}
    - 课程图片信息: 见附带图片
    【内容要求】
    1. 多模态整合：综合解析提供的文本列表与图片信息，提取课程核心关键词。
    2. 精准筛选：仅保留与申请目标强相关的课程；若为构建逻辑递进感所必需，可保留关键基础课。
    3. 动态知识补全，若缺课程说明，需主动检索硕士级别的教学大纲，检索重点是该课程的核心方法论（Methodology）与前沿概念（Concepts）。
    4. 深度叙事构建：
        非线性罗列：严禁简单重复课程名。
        逻辑升华：按“知识难度递进”或“学科内在逻辑”排列，体现从理论基础到复杂应用的认知过程。
        学术还原：详细阐述具体的方法学、模型或理论，并解释其对学生的学术吸引力与实际帮助。
    5. 语调约束：保持朴素、专业的学术笔调，以议论与反思为主（Argumentative & Reflective），拒绝浮夸口吻。
    {CLEAN_OUTPUT_RULES}
    """

    prompt_internship = f"""
    【任务】撰写 "实习/工作经历" (Professional Experience) 部分。
    【输入背景】
    - 学生素材: 见附带文本
    - 目标专业: {target_school_name}
    【内容要求】
    1. 精选叙事：仅挑选与目标专业最强相关的经历。严禁全量罗列，确保“质量”凌驾于“数量”
    2. 线性逻辑：按时间先后顺序串联选定的经历，构建个人成长的进化路径。
    3. 要素融合 (B-R-S-M)：每段经历须包含以下维度，但需自然揉合，拒绝机械填表：
        Context (背景)：任务的起因或环境。
        Responsibility (职责)：你承担的具体角色。
        Skills (技能)：所运用的核心硬核能力。
        Motivation (动机)：驱动你行动的内在逻辑及事后感悟。
    4. 深度反思：
        拒绝平铺直叙的任务描述（流水账）。
        必须挖掘每项经历背后的学术/专业感悟。
        明确指出该经历如何体现你与目标专业的契合度（Match Point）。
    5. 语调要求：保持职场与学术兼具的专业笔调，以逻辑论证为叙述核心。
    {CLEAN_OUTPUT_RULES}
    """

    prompts_map = {
        "Motivation": prompt_motivation,
        "Career_Goal": prompt_career,
        "Academic": prompt_academic,
        "Why_School": prompt_whyschool,
        "Internship": prompt_internship
    }

    for module in selected_modules:
        current_step += 1
        
        status_container.info(f"正在撰写: {modules[module]} ({current_step}/{total_steps})")
        detail_container.markdown(f"**分析背景资料**，构建 {modules[module]} 部分...")
        
        current_media = None
        if module == "Academic":
            current_media = transcript_content
            detail_container.markdown("**分析成绩单**，提取相关学术背景...")
        elif module == "Why_School":
            current_media = curriculum_imgs
            detail_container.markdown("**分析课程信息**，匹配学生背景与课程优势...")
        
        detail_container.markdown("**撰写初稿中**，请稍候...")
        
        res = get_gemini_response(prompts_map[module], media_content=current_media, text_context=student_background_text)
        
        final_text = res.strip()
        
        # 稳健性优化：处理解析失败的情况
        if module == "Motivation":
            try:
                if "[TRENDS_START]" in res and "[DRAFT_START]" in res:
                    trends_part = res.split("[TRENDS_START]")[1].split("[TRENDS_END]")[0].strip()
                    draft_part = res.split("[DRAFT_START]")[1].split("[DRAFT_END]")[0].strip()
                    st.session_state['motivation_trends'] = trends_part
                    final_text = draft_part
                    detail_container.markdown("提取行业趋势，整合到申请动机...")
                else:
                    final_text = res
            except Exception as e:
                # 兜底逻辑：如果分割出错，直接保留全文
                final_text = res

        st.session_state['generated_sections'][module] = final_text
        progress_bar.progress(current_step / total_steps)
        detail_container.markdown(f"**{modules[module]}** 部分已完成！")
        time.sleep(0.5)

    detail_container.markdown("整合所有部分，生成完整草稿...")
    full_chinese_draft = ""
    for module in display_order:
        if module in st.session_state['generated_sections']:
            full_chinese_draft += f"--- {modules[module]} ---\n"
            full_chinese_draft += st.session_state['generated_sections'][module] + "\n\n"
    st.session_state['full_chinese_draft'] = full_chinese_draft.strip()
    
    st.session_state['full_translated_text'] = ""
    if 'text_full_draft' in st.session_state:
        del st.session_state['text_full_draft']
    if 'text_full_translated' in st.session_state:
        del st.session_state['text_full_translated']
    
    if 'header_cn' in st.session_state:
        del st.session_state['header_cn']
    if 'header_en' in st.session_state:
        del st.session_state['header_en']

    status_container.empty()
    detail_container.empty()

    st.markdown(f"""
    <div style="
        background-color: #3666FA; 
        color: #ffffff; 
        padding: 15px; 
        border-radius: 12px; 
        text-align: center; 
        font-weight: 600;
        margin-top: 20px;
        box-shadow: 0 2px 5px rgba(54, 102, 250, 0.2);
    ">
        初稿生成完毕！
    </div>
    """, unsafe_allow_html=True)

# ==========================================
# 8. 界面：反馈、修改与翻译
# ==========================================
if st.session_state.get('full_chinese_draft'):
    st.markdown("---")
    st.header("审阅与翻译")
    
    render_blue_box("满意左侧中文稿后，点击翻译按钮生成翻译。")

    if st.session_state.get('motivation_trends'):
        with st.expander("点击查看：行业趋势调研与参考源 (Reference)", expanded=True):
            render_blue_box(st.session_state['motivation_trends'])
    
    style_text = "British" if "British" in spelling_preference else "American"
    
    button_row = st.columns([1, 1, 1])
    with button_row[0]:
        translate_button = st.button(f"翻译全文 ({style_text})", key="translate_btn")
    
    with button_row[2]:
        if st.session_state.get('full_translated_text'):
            english_edit_button = st.button("执行英文批注修改", key="english_edit_btn")
    
    c1, c2 = st.columns(2)
    
    # --- 左侧：中文编辑 ---
    with c1:
        st.markdown("**中文草稿 (可编辑)**")
        
        if 'text_full_draft' not in st.session_state:
            st.session_state['text_full_draft'] = st.session_state['full_chinese_draft']
        
        current_chinese_content = st.text_area(
            "中文内容", 
            key="text_full_draft",
            height=600
        )
        st.session_state['full_chinese_draft'] = current_chinese_content
        
        render_blue_box("批注修改: 在想改的句子后面用 【修改意见】 给出指令。")
        
        if st.button("执行中文批注修改", key="chinese_edit_btn"):
            if "【" not in current_chinese_content:
                st.warning("未检测到【】。请在上方文本框中插入 `【修改意见】` 后再点击。")
            else:
                with st.spinner("正在根据批注修改并高亮变化..."):
                    inline_prompt = f"""
                    【任务】作为专业留学文书编辑，根据文中的嵌入式批注（中文方括号【】内的文字）修改文章。
                    【输入文本】\n{current_chinese_content}
                    【执行步骤】
                    1. 扫描文中所有的中文方括号 `【】`。括号内的文字即为用户的修改指令。
                    2. 根据指令，修改括号紧邻的前文句子或段落。
                    3. **必须删除**原文中的括号及括号内的修改指令。
                    4. 保持未被批注的部分原封不动。
                    5. **高亮变化**：将**所有被修改后产生的新文字**用 Markdown 双星号 `**` 包裹（例如：**new text**），以便用户一眼看出改了哪里。
                    {CLEAN_OUTPUT_RULES}
                    """
                    revised_text = get_gemini_response(inline_prompt)
                    
                    st.session_state['full_chinese_draft'] = revised_text.strip()
                    if 'text_full_draft' in st.session_state:
                        del st.session_state['text_full_draft'] 
                    st.session_state['full_translated_text'] = ""
                    if 'text_full_translated' in st.session_state:
                        del st.session_state['text_full_translated']
                    st.rerun()

    # --- 右侧：英文翻译 ---
    with c2:
        st.markdown("**英文翻译结果 (可编辑)**")
        
        if st.session_state.get('full_translated_text'):
            if 'text_full_translated' not in st.session_state:
                st.session_state['text_full_translated'] = st.session_state['full_translated_text']

            current_english_content = st.text_area(
                "英文内容",
                key="text_full_translated",
                height=600
            )
            st.session_state['full_translated_text'] = current_english_content

            render_blue_box("批注修改: 在想改的句子后面用 【修改意见】 给出指令。")
        else:
            st.text_area(
                "等待翻译...",
                value="点击上方的翻译按钮生成英文翻译。",
                height=600,
                disabled=True
            )
            render_blue_box("满意左侧中文稿后，点击上方按钮生成翻译。")
    
    # 翻译逻辑
    if translate_button:
        if not api_key:
            st.error("需要 API Key")
        else:
            with st.spinner("Translating..."):
                spelling_instruction = "\n【SPELLING RULE】: STRICTLY use British English spelling (e.g., colour, analyse, programme, centre)."
                if "American" in spelling_preference:
                    spelling_instruction = "\n【SPELLING RULE】: STRICTLY use American English spelling (e.g., color, analyze, program, center)."
                
                translated_sections = []
                for module_key in display_order:
                    if module_key in st.session_state['generated_sections']:
                        chinese_text = st.session_state['generated_sections'][module_key]
                        if not chinese_text.strip():
                            continue
                        
                        trans_prompt = f"{TRANSLATION_RULES_BASE}\n{spelling_instruction}\n【Input Text】:\n{chinese_text}"
                        trans_res = get_gemini_response(trans_prompt)

                        english_header = english_modules.get(module_key, module_key)
                        translated_sections.append(f"--- {english_header} ---\n{trans_res.strip()}")
                
                st.session_state['full_translated_text'] = "\n\n".join(translated_sections)
                
                if 'text_full_translated' in st.session_state:
                    del st.session_state['text_full_translated']
                st.rerun()
    
    # 英文修改逻辑
    if st.session_state.get('full_translated_text') and 'english_edit_button' in locals() and english_edit_button:
        with st.spinner("正在根据您的批注优化英文文本..."):
            current_english_content = st.session_state['full_translated_text']
            english_edit_prompt = f"""
            【任务】你是一位顶尖的留学文书编辑。请根据用户在英文文本中嵌入的中文，对文章进行修改和润色。

            【输入文本及批注】
            {current_english_content}

            【批注规则说明】
            1.  **修改指令 `【中文内容】`**: 如果发现中文被中文方括号 `【】` 包围，这代表一条修改指令。请根据指令内容，修改它前面的英文句子。
            2.  **翻译并插入**: 如果发现一段中文**没有被任何括号包围**，请将这段中文翻译成地道的英文，并无缝地插入到文本的那个位置。

            【核心风格指令】
            所有的修改和翻译都必须严格遵守以下【ANTI-AI STYLE GUIDE】。
            {TRANSLATION_RULES_BASE}

            【输出要求】
            1.  完成所有修改和翻译。
            2.  **必须删除**原文中所有的中文内容和 `【】` 括号。
            3.  **必须保留**所有的分段标题（例如 `--- Motivation ---`）。
            4.  将**所有被修改或新增的英文部分**用 Markdown 双星号 `**` 包裹，以便用户识别。
            5.  最终输出完整的、保留了分段结构的英文文本。
            """
            revised_english_text = get_gemini_response(english_edit_prompt)
            st.session_state['full_translated_text'] = revised_english_text.strip()
            if 'text_full_translated' in st.session_state:
                del st.session_state['text_full_translated']
            st.rerun()

# ==========================================
# 9. 导出
# ==========================================
if st.session_state.get('full_chinese_draft'):
    st.markdown("---")
    st.header("导出")
    
    if 'header_cn' not in st.session_state or 'header_en' not in st.session_state:
        if target_school_name:
            header_prompt = f"""
            Task: Parse and format the university and major information from the string: "{target_school_name}".
            
            Rules:
            1. Identify the School Name and Major Name.
            2. Create a Chinese Header: [School Name (Chinese, add '大学' if missing)] + [Major Name] + "个人陈述"
            3. Create an English Header: "Personal Statement for " + [Major Name (English)] + "_" + [School Name (English)]
            
            Example Input: 卡内基梅隆Master's in Health Care Analytics
            Example Output: 卡内基梅隆大学Master's in Health Care Analytics个人陈述|Personal Statement for Master's in Health Care Analytics_Carnegie Mellon University
            
            Output ONLY the two strings separated by a pipe symbol (|). Do not add any other text.
            """
            try:
                header_res = get_gemini_response(header_prompt)
                if "|" in header_res:
                    parts = header_res.split("|")
                    st.session_state['header_cn'] = parts[0].strip()
                    st.session_state['header_en'] = parts[1].strip()
                else:
                    st.session_state['header_cn'] = f"{target_school_name} 个人陈述"
                    st.session_state['header_en'] = f"Personal Statement for {target_school_name}"
            except:
                st.session_state['header_cn'] = f"{target_school_name} 个人陈述"
                st.session_state['header_en'] = f"Personal Statement for {target_school_name}"
        else:
             st.session_state['header_cn'] = "个人陈述"
             st.session_state['header_en'] = "Personal Statement"

    col_dl_cn, col_dl_en = st.columns(2)
    
    with col_dl_cn:
        st.subheader("中文版")
        if st.session_state.get('full_chinese_draft'):
            cn_header_text = st.session_state.get('header_cn', f"{target_school_name} 个人陈述")
            docx_cn_bytes = create_word_docx(
                content=st.session_state['full_chinese_draft'],
                header_text=cn_header_text,
                font_name='宋体',
                is_chinese=True
            )
            
            st.download_button(
                label="下载中文版 (.docx)",
                data=docx_cn_bytes,
                file_name=f"PS_CN_{target_school_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )
        else:
            st.caption("暂无中文内容")

    with col_dl_en:
        st.subheader("英文版")
        if st.session_state.get('full_translated_text'):
            en_header_text = st.session_state.get('header_en', f"Personal Statement for {target_school_name}")
            docx_en_bytes = create_word_docx(
                content=st.session_state['full_translated_text'],
                header_text=en_header_text,
                font_name='Times New Roman',
                is_chinese=False
            )
            
            st.download_button(
                label="下载英文版 (.docx)",
                data=docx_en_bytes,
                file_name=f"PS_EN_{target_school_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )
        else:
            st.caption("暂无英文翻译，请先在上方进行翻译。")
