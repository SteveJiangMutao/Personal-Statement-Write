from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
import google.generativeai as genai
from PIL import Image
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pypdf
import io
import os
import time
import random
import re
from datetime import datetime
from typing import List, Optional, Dict, Any
import json
from pydantic import BaseModel
import base64

app = FastAPI(title="Personal Statement Writing API", version="1.0.0")

# Environment variables
GOOGLE_API_KEY = os.environ.get('GOOGLE_API_KEY')

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, replace with frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==========================================
# 1. 数据模型
# ==========================================
class GenerationRequest(BaseModel):
    api_key: str
    model_name: str = "gemini-2.5-pro"
    target_school_name: str
    counselor_strategy: str = ""
    selected_modules: List[str]
    spelling_preference: str = "British"  # "British" or "American"
    material_text: Optional[str] = None
    # Files will be handled separately as multipart form data

class FileUploadRequest(BaseModel):
    api_key: str
    model_name: str = "gemini-2.5-pro"
    target_school_name: str
    counselor_strategy: str = ""
    selected_modules: List[str]
    spelling_preference: str = "British"
    material_file: Optional[UploadFile] = None
    transcript_file: Optional[UploadFile] = None
    curriculum_text: Optional[str] = None
    curriculum_files: Optional[List[UploadFile]] = None

class TranslationRequest(BaseModel):
    api_key: str = ""
    model_name: str = "gemini-2.5-pro"
    chinese_text: str
    spelling_preference: str = "British"
    module_type: str  # "Motivation", "Academic", etc.

class EditRequest(BaseModel):
    api_key: str = ""
    model_name: str = "gemini-2.5-pro"
    text: str
    is_chinese: bool = True

class WordGenerationRequest(BaseModel):
    content: str
    header_text: str
    is_chinese: bool = False
    font_name: str = "宋体"

# ==========================================
# 2. 核心辅助函数 (从原 psw.py 移植)
# ==========================================
def set_bottom_border(paragraph):
    """为段落添加下框线 (用于页眉)"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6') # 1/8 pt, 6 = 0.75pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000') # 黑色
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_word_docx(content, header_text, font_name, is_chinese=False):
    """生成 Word 文档 (包含清洗逻辑)"""
    doc = docx.Document()

    # --- 1. 设置页眉 ---
    section = doc.sections[0]
    header = section.header

    # 获取页眉的第一个段落（默认存在）
    header_para = header.paragraphs[0]
    header_para.text = header_text
    header_para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

    # 设置页眉下框线
    set_bottom_border(header_para)

    # 设置页眉字体样式 (12pt, 斜体)
    for run in header_para.runs:
        run.font.name = font_name
        run.font.size = Pt(12)
        run.font.italic = True
        # 处理中文字体显示
        if is_chinese:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # --- 2. 设置正文 (清洗逻辑优化) ---
    # 1. 去除 Markdown 加粗符号
    content = content.replace("**", "")
    # 2. 去除 Markdown 单星号 (列表或斜体)
    content = content.replace("*", "")

    # 按行处理
    for line in content.split('\n'):
        line = line.strip()

        # 3. 跳过空行
        if not line:
            continue

        # 4. 🚨 核心修改：跳过段落标题行 (特征：以 --- 开头)
        # 确保只保留正文，移除类似 "--- Motivation ---" 或 "--- 申请动机 ---" 的行
        if line.startswith("---") and line.endswith("---"):
            continue

        p = doc.add_paragraph(line)
        # 设置正文样式 (11pt)
        for run in p.runs:
            run.font.name = font_name
            run.font.size = Pt(11)
            # 处理中文字体显示
            if is_chinese:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # 保存到内存
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def read_word_file(file_bytes):
    """读取 Word 文件内容"""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading Word file: {e}"

def read_pdf_text(file_bytes):
    """读取 PDF 文件内容"""
    try:
        pdf_reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Error reading PDF file: {e}"

def get_gemini_response(api_key: str, model_name: str, prompt: str, media_content=None, text_context=None):
    """调用 Gemini API"""
    # 优先使用环境变量中的API Key
    effective_api_key = GOOGLE_API_KEY if GOOGLE_API_KEY else api_key
    if not effective_api_key:
        return "Error: API Key is required. Please set GOOGLE_API_KEY environment variable or provide via request."

    genai.configure(api_key=effective_api_key)
    model = genai.GenerativeModel(model_name)

    content = []
    content.append(prompt)

    if text_context:
        content.append(f"\n【参考文档/背景信息 (简历或素材表)】:\n{text_context}")

    if media_content:
        if isinstance(media_content, list):
            content.extend(media_content)
        else:
            content.append(media_content)

    try:
        response = model.generate_content(content)
        return response.text
    except Exception as e:
        return f"Error: {str(e)}"

# ==========================================
# 3. 提示词模板 (从原 psw.py 移植)
# ==========================================
CLEAN_OUTPUT_RULES = """
【绝对输出规则】
1. 只输出正文内容本身。
2. 严禁包含开场白、结尾语或结构说明。
3. 严禁使用 Markdown 格式（如加粗、列表符号、标题符号）。
4. 输出必须是纯文本。
5. 必须写成一个完整的、连贯的中文自然段。
"""

TRANSLATION_RULES_BASE = """
【Translation Task】
Translate the provided Chinese text into a professional, human-sounding Personal Statement paragraph.

【CRITICAL ANTI-AI STYLE GUIDE】
1. **KILL THE "AI SENTENCE PATTERN"**:
   - **ABSOLUTELY FORBIDDEN**: The pattern "I did X, **thereby/thus/enabling** me to do Y."
   - **SOLUTION**: Split into two sentences or use active verbs.

2. **SEMICOLONS (;) FOR FLOW**:
   - **MANDATORY**: When a sentence is grammatically complete but the thought is not finished (and leads directly into the next point), use a **semicolon (;)** to connect them.

3. **ADVERB CONTROL (ZERO TOLERANCE)**:
   - **STRICTLY PROHIBITED**: The combination of **Adverb + Verb** (e.g., "deeply analyze", "successfully completed") OR **Adverb + Adjective** (e.g., "perfectly align", "keenly interested").
   - **ACTION**: Delete the adverb entirely. Just use the verb or adjective.

4. **VOCABULARY PURGE**:
   - Use precise, simple words.

5. **ENHANCE COHESION & NARRATIVE FLOW (CRITICAL)**:
   - **MANDATORY**: You MUST actively add varied transitional phrases and logical connectors (e.g., "Furthermore," "In contrast," "Consequently," "Given this context") between sentences AND between paragraphs.
   - **GOAL**: Ensure the text flows smoothly as a unified narrative, not a disjointed list of sentences. The priority is reading fluency and the overall integrity of the article.

【BANNED WORDS LIST (Strictly Prohibited)】
[Verbs]: delve into, uncover, reveal, recognize, master, refine, cultivate, address, bridge, spearhead, pioneer, align with, stems from, underscore, highlight
[Adjectives/Adverbs]: instrumental, pivotal, seamless, systematically, rigorously, profoundly, deeply, acutely, keenly, comprehensively, perfectly, meticulously, proficiency, Additionally
[Nouns]: paradigm, trajectory, aspirations, vision, landscape, tapestry, realm, foundation, tenure, testament, commitment
[Connectors]: thereby, thus (when used with -ing), in turn
[Phrases]: "not only... but also", "Building on this", "rich tapestry", "testament to", "a wide array of", "my goal is to"， "focus will be"

【Formatting】
1. Output as ONE single paragraph.
2. Output the ENTIRE text in **Bold**.
3. No Markdown headers.
"""

def get_prompt_motivation(target_school_name: str) -> str:
    return f"""
    【任务】撰写 Personal Statement 的 "申请动机" 部分。
    【步骤 1：深度调研】
    请先分析 {target_school_name} 所在领域的最新行业热点或学术趋势。
    **请严格列出 3 个关键趋势 (Options)**，并严格按照以下 **HTML 格式** 输出（除文献/报告标题保留原文外，其余分析内容请使用**中文**）：

    <div style="margin-bottom: 18px;">
        <div style="font-weight: bold; font-size: 14px; margin-bottom: 6px;">Option [X]: [Trend Title]</div>
        <ul style="margin: 0; padding-left: 18px; list-style-position: outside;">
            <li style="margin-bottom: 4px; line-height: 1.4;"><b>Source</b>: [Specific Paper Title/Report Name/News Source]</li>
            <li style="line-height: 1.4;"><b>Relevance</b>: [深度分析趋势与学生背景/项目的关联。解释为什么这个趋势对该学生重要，以及他们之前的经历（如具体项目、技能）如何与此契合。此部分必须详细展开。]</li>
        </ul>
    </div>

    【步骤 2：撰写正文】
    基于上述趋势和学生素材，撰写一段中文申请动机。动机正文中不用出现具体信息源，但要体现出学生对行业趋势的理解和契合。
    逻辑：学生过往经历 -> 观察到的行业痛点/趋势 -> 产生深造需求。
    【严格输出格式】
    请严格按照下方分隔符输出，不要包含其他内容：
    [TRENDS_START]
    (在此处列出 3 个调研趋势和来源，使用上述 HTML 格式)
    [TRENDS_END]
    [DRAFT_START]
    (在此处撰写正文段落，纯文本，无Markdown)
    [DRAFT_END]
    """

def get_prompt_career(target_school_name: str, counselor_strategy: str) -> str:
    return f"""
    【任务】撰写 "职业规划" (Career Goals) 部分。
    【输入背景】
    - 目标专业: {target_school_name}
    - 顾问思路: {counselor_strategy}
    【内容要求】
    1. 规划硕士毕业后的路径（应届生视角）。
    2. **必须包含**：具体的公司名字、具体的职位名称。
    3. 将工作内容和未来继续学习方向融合在一段话中。
    {CLEAN_OUTPUT_RULES}
    """

def get_prompt_academic(target_school_name: str) -> str:
    return f"""
    【任务】撰写 "本科学习经历" (Academic Background) 部分。
    【输入背景】
    - 目标专业: {target_school_name}
    - 核心依据 (成绩单): 见附带文件 (PDF或图片)
    - 辅助参考 (学生素材/简历): 见附带文本
    【核心原则：深度 > 数量】
    不要罗列课程名。只精选与目标专业最强相关的核心课程进行深度描写。
    【内容要求 - 必须包含细节】
    1. **核心概念植入**：在描述每门课时，必须提及该课程具体的**核心概念、模型、算法或理论名称**。
    2. **学术真实感**：结合学生素材，简述是如何理解或应用这些概念的。
    3. **逻辑升华**：说明这些具体的知识点如何为你攻读 {target_school_name} 打下了坚实的学术基础。
    4. **禁止**：禁止写成课程清单（List），必须是连贯的学术反思叙述。
    {CLEAN_OUTPUT_RULES}
    """

def get_prompt_whyschool(target_school_name: str, counselor_strategy: str, target_curriculum_text: str) -> str:
    return f"""
    【任务】撰写 "Why School" 部分。
    【输入背景】
    - 目标学校: {target_school_name}
    - 顾问思路: {counselor_strategy}
    {f'【目标课程文本列表】:{target_curriculum_text}' if target_curriculum_text else ''}
    - 课程图片信息: 见附带图片
    【内容要求】
    1. 综合分析提供的文本列表和图片中的课程信息。
    2. 从中挑选与学生背景或规划最相关的特定课程，不相关的课程不用写。
    3. 若所提供信息包含课程名字与课程说明则参考，若仅有课程名字但无课程说明则搜索该课程（硕士水平）的教学内容，并据此阐述这些课程为何吸引学生及有何帮助，阐述时需深入到该课程具体教授的方法学及概念。
    4. 课程阐述需有深度，有逻辑顺序或难度递进关系，体现出对课程内容的理解，而非简单罗列课程名称。
    5. 语气朴素专业，议论为主。
    {CLEAN_OUTPUT_RULES}
    """

def get_prompt_internship(target_school_name: str) -> str:
    return f"""
    【任务】撰写 "实习/工作经历" (Professional Experience) 部分。
    【输入背景】
    - 学生素材: 见附带文本
    - 目标专业: {target_school_name}
    【内容要求】
    1. 筛选最相关经历，按时间顺序逻辑串联。
    2. 结构：背景 -> 职责 -> 技能 -> 动机。
    3. 拒绝流水账，要有逻辑梳理和反思，要有与所申请专业的契合点和相关的感悟。
    {CLEAN_OUTPUT_RULES}
    """

# 模块映射
modules = {
    "Motivation": "申请动机",
    "Academic": "本科学习",
    "Internship": "实习/工作",
    "Why_School": "选校理由",
    "Career_Goal": "职业规划"
}

english_modules = {
    "Motivation": "Motivation",
    "Academic": "Academic Background",
    "Internship": "Professional Experience",
    "Why_School": "Why School",
    "Career_Goal": "Career Goal"
}

display_order = ["Motivation", "Academic", "Internship", "Why_School", "Career_Goal"]

# ==========================================
# 4. API 端点
# ==========================================
@app.get("/")
def read_root():
    return {"message": "Personal Statement Writing API", "status": "running"}

@app.post("/api/generate")
async def generate_personal_statement(
    api_key: str = Form(""),
    model_name: str = Form("gemini-2.5-pro"),
    target_school_name: str = Form(...),
    counselor_strategy: str = Form(""),
    selected_modules: str = Form(...),  # JSON string of list
    spelling_preference: str = Form("British"),
    material_file: Optional[UploadFile] = File(None),
    transcript_file: Optional[UploadFile] = File(None),
    curriculum_text: Optional[str] = Form(None),
    curriculum_files: Optional[List[UploadFile]] = File([]),
):
    """生成个人陈述各个模块的内容"""
    try:
        # Parse selected modules
        modules_list = json.loads(selected_modules)

        # Read material file
        student_background_text = ""
        if material_file:
            file_bytes = await material_file.read()
            if material_file.filename.endswith('.docx'):
                student_background_text = read_word_file(file_bytes)
            elif material_file.filename.endswith('.pdf'):
                student_background_text = read_pdf_text(file_bytes)

        # Prepare media content
        transcript_content = []
        if transcript_file:
            file_bytes = await transcript_file.read()
            if transcript_file.content_type == "application/pdf":
                transcript_content.append({
                    "mime_type": "application/pdf",
                    "data": file_bytes
                })
            else:
                # For image files
                transcript_content.append(Image.open(io.BytesIO(file_bytes)))

        curriculum_imgs = []
        if curriculum_files:
            for img_file in curriculum_files:
                file_bytes = await img_file.read()
                curriculum_imgs.append(Image.open(io.BytesIO(file_bytes)))

        # Generate content for each selected module
        generated_sections = {}
        motivation_trends = ""

        for module in modules_list:
            # Get appropriate prompt
            if module == "Motivation":
                prompt = get_prompt_motivation(target_school_name)
                current_media = None
            elif module == "Career_Goal":
                prompt = get_prompt_career(target_school_name, counselor_strategy)
                current_media = None
            elif module == "Academic":
                prompt = get_prompt_academic(target_school_name)
                current_media = transcript_content
            elif module == "Why_School":
                prompt = get_prompt_whyschool(target_school_name, counselor_strategy, curriculum_text or "")
                current_media = curriculum_imgs
            elif module == "Internship":
                prompt = get_prompt_internship(target_school_name)
                current_media = None
            else:
                continue

            # Call Gemini API
            response = get_gemini_response(
                api_key=api_key,
                model_name=model_name,
                prompt=prompt,
                media_content=current_media,
                text_context=student_background_text
            )

            final_text = response.strip()

            # Special handling for Motivation module
            if module == "Motivation":
                if "[TRENDS_START]" in response and "[DRAFT_START]" in response:
                    trends_part = response.split("[TRENDS_START]")[1].split("[TRENDS_END]")[0].strip()
                    draft_part = response.split("[DRAFT_START]")[1].split("[DRAFT_END]")[0].strip()
                    motivation_trends = trends_part
                    final_text = draft_part
                else:
                    final_text = response

            generated_sections[module] = final_text

        # Build full Chinese draft
        full_chinese_draft = ""
        for module in display_order:
            if module in generated_sections:
                full_chinese_draft += f"--- {modules[module]} ---\n"
                full_chinese_draft += generated_sections[module] + "\n\n"

        return JSONResponse(content={
            "success": True,
            "generated_sections": generated_sections,
            "full_chinese_draft": full_chinese_draft.strip(),
            "motivation_trends": motivation_trends
        })

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation failed: {str(e)}")

@app.post("/api/translate")
async def translate_content(request: TranslationRequest):
    """翻译中文内容到英文"""
    try:
        spelling_instruction = "\n【SPELLING RULE】: STRICTLY use British English spelling (e.g., colour, analyse, programme, centre)."
        if request.spelling_preference == "American":
            spelling_instruction = "\n【SPELLING RULE】: STRICTLY use American English spelling (e.g., color, analyze, program, center)."

        trans_prompt = f"{TRANSLATION_RULES_BASE}\n{spelling_instruction}\n【Input Text】:\n{request.chinese_text}"

        translated_text = get_gemini_response(
            api_key=request.api_key,
            model_name=request.model_name,
            prompt=trans_prompt
        )

        return JSONResponse(content={
            "success": True,
            "translated_text": translated_text.strip(),
            "module": request.module_type
        })

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")

@app.post("/api/edit")
async def edit_content(request: EditRequest):
    """根据批注编辑内容"""
    try:
        if request.is_chinese:
            inline_prompt = f"""
            【任务】作为专业留学文书编辑，根据文中的嵌入式批注（中文方括号【】内的文字）修改文章。
            【输入文本】\n{request.text}
            【执行步骤】
            1. 扫描文中所有的中文方括号 `【】`。括号内的文字即为用户的修改指令。
            2. 根据指令，修改括号紧邻的前文句子或段落。
            3. **必须删除**原文中的括号及括号内的修改指令。
            4. 保持未被批注的部分原封不动。
            5. **高亮变化**：将**所有被修改后产生的新文字**用 Markdown 双星号 `**` 包裹（例如：**new text**），以便用户一眼看出改了哪里。
            {CLEAN_OUTPUT_RULES}
            """
        else:
            inline_prompt = f"""
            【任务】你是一位顶尖的留学文书编辑。请根据用户在英文文本中嵌入的中文，对文章进行修改和润色。

            【输入文本及批注】
            {request.text}

            【批注规则说明】
            1.  **修改指令 `【中文内容】`**: 如果发现中文被中文方括号 `【】` 包围，这代表一条修改指令。请根据指令内容，修改它前面的英文句子。
            2.  **翻译并插入**: 如果发现一段中文**没有被任何括号包围**，请将这段中文翻译成地道的英文，并无缝地插入到文本的那个位置。

            【核心风格指令】
            所有的修改和翻译都必须严格遵守以下【ANTI-AI STYLE GUIDE】。
            {TRANSLATION_RULES_BASE}

            【输出要求】
            1.  完成所有修改和翻译。
            2.  **必须删除**原文中所有的中文内容和 `【】` 括号。
            3.  **必须保留**所有的分段标题（例如 `--- Motivation ---`）。
            4.  将**所有被修改或新增的英文部分**用 Markdown 双星号 `**` 包裹，以便用户识别。
            5.  最终输出完整的、保留了分段结构的英文文本。
            """

        edited_text = get_gemini_response(
            api_key=request.api_key,
            model_name=request.model_name,
            prompt=inline_prompt
        )

        return JSONResponse(content={
            "success": True,
            "edited_text": edited_text.strip()
        })

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Edit failed: {str(e)}")

@app.post("/api/generate-word")
async def generate_word_document(request: WordGenerationRequest):
    """生成Word文档"""
    try:
        docx_bytes = create_word_docx(
            content=request.content,
            header_text=request.header_text,
            font_name=request.font_name,
            is_chinese=request.is_chinese
        )

        # Determine filename
        if request.is_chinese:
            filename = "personal_statement_cn.docx"
        else:
            filename = "personal_statement_en.docx"

        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Word generation failed: {str(e)}")

@app.post("/api/generate-header")
async def generate_header(
    api_key: str = Form(""),
    model_name: str = Form("gemini-2.5-pro"),
    target_school_name: str = Form(...)
):
    """生成中英文页眉"""
    try:
        header_prompt = f"""
        Task: Parse and format the university and major information from the string: "{target_school_name}".

        Rules:
        1. Identify the School Name and Major Name.
        2. Create a Chinese Header: [School Name (Chinese, add '大学' if missing)] + [Major Name] + "个人陈述"
        3. Create an English Header: "Personal Statement for " + [Major Name (English)] + "_" + [School Name (English)]

        Example Input: 卡内基梅隆Master's in Health Care Analytics
        Example Output: 卡内基梅隆大学Master's in Health Care Analytics个人陈述|Personal Statement for Master's in Health Care Analytics_Carnegie Mellon University

        Output ONLY the two strings separated by a pipe symbol (|). Do not add any other text.
        """

        header_res = get_gemini_response(
            api_key=api_key,
            model_name=model_name,
            prompt=header_prompt
        )

        if "|" in header_res:
            parts = header_res.split("|")
            header_cn = parts[0].strip()
            header_en = parts[1].strip()
        else:
            # Fallback
            header_cn = f"{target_school_name} 个人陈述"
            header_en = f"Personal Statement for {target_school_name}"

        return JSONResponse(content={
            "success": True,
            "header_cn": header_cn,
            "header_en": header_en
        })

    except Exception as e:
        return JSONResponse(content={
            "success": True,
            "header_cn": f"{target_school_name} 个人陈述",
            "header_en": f"Personal Statement for {target_school_name}"
        })

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)